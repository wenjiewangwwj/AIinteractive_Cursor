"""Turn Streamlit uploads into text context and multimodal image parts."""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Any

from docx import Document
from pypdf import PdfReader


MAX_TEXT_CONTEXT_CHARS = 120_000


@dataclass
class PreparedContext:
    """Unified attachment payload for both providers."""

    text_block: str
    images: list[tuple[str, bytes]]  # (media_type e.g. image/png, raw bytes)


def _truncate(s: str) -> str:
    if len(s) <= MAX_TEXT_CONTEXT_CHARS:
        return s
    return s[: MAX_TEXT_CONTEXT_CHARS] + "\n\n[Truncated: attachment text exceeded limit.]"


def _pdf_to_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text() or ""
        parts.append(t)
    return "\n\n".join(parts).strip() or "[PDF: no extractable text]"


def _docx_to_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip() or "[DOCX: no extractable text]"


def prepare_uploads(files: list[Any] | None) -> PreparedContext:
    text_chunks: list[str] = []
    images: list[tuple[str, bytes]] = []

    if not files:
        return PreparedContext(text_block="", images=[])

    for f in files:
        name = getattr(f, "name", "upload") or "upload"
        data = f.getvalue()
        lower = name.lower()

        image_types = (
            (".png", "image/png"),
            (".jpg", "image/jpeg"),
            (".jpeg", "image/jpeg"),
            (".gif", "image/gif"),
            (".webp", "image/webp"),
        )
        matched = next((mt for suf, mt in image_types if lower.endswith(suf)), None)
        if matched:
            images.append((matched, data))
            text_chunks.append(f"[Attached image: {name}]\n")
        elif lower.endswith(".pdf"):
            text_chunks.append(f"--- PDF: {name} ---\n{_pdf_to_text(data)}\n")
        elif lower.endswith(".docx"):
            try:
                text_chunks.append(f"--- DOCX: {name} ---\n{_docx_to_text(data)}\n")
            except Exception as e:  # noqa: BLE001
                text_chunks.append(
                    f"--- DOCX: {name} ---\n[Could not read DOCX: {e}]\n"
                )
        else:
            try:
                text_chunks.append(f"--- File: {name} ---\n{data.decode('utf-8')}\n")
            except UnicodeDecodeError:
                text_chunks.append(
                    f"--- File: {name} ---\n[Skipped: not UTF-8 text. Use PDF, DOCX, or image.]\n"
                )

    combined = "\n".join(text_chunks).strip()
    return PreparedContext(text_block=_truncate(combined), images=images)
